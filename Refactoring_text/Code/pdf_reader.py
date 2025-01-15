import fitz
import re
import spacy
from spacy.symbols import ORTH, LEMMA
from docx import Document
import pymorphy2


# Цвета для вывода
class bcolors:
    OKGREEN = '\033[92m'
    YELLOW = '\033[93m'
    ENDC = '\033[0m'
    UND = '\033[4m'
    BOLD = '\033[1m'
    CYAN = '\033[96m' # {bcolors.CYAN}{....}{bcolors.ENDC}


pdf_file = ["Положение о Минобороны.pdf",
            "ПМО 2014 г. № 400дсп Положение о ГОУ.pdf",
            "ПМО 2016 г. № 260дсп Положение о ГК СВ.pdf",
            "Приказ Министр обороны Российской Федерации от 16.10.2018 № 580 (в редакции от 27.01.2024 № 41).pdf"]

pdf_path = "docs/" + pdf_file[1]
text = ''
file = fitz.open(pdf_path)
for pageNum, page in enumerate(file.pages()):
    text += page.get_text()

# Создаётся строковая переменная из всех данных
result = text.replace('\n', ' ')
# Удаление номеров страниц
result = re.sub(r'( {2,}[0-9]{1,2} {2,})', ' ', result)
# Удаляются лишние пробелы
result = re.sub(r' {2,}', ' ', result)

# Удаление текста в скобках о редакциях, пунктах и т.д.
result = re.sub(r'(\(в ред\. [а-яА-Я0-9.,;№N ]+\) |\(п\. [0-9]*\.?[0-9]* [а-яА-Я0-9.,;№N ]+\) |\(пп\. [0-9]*\.?[0-9]* [а-яА-Я0-9.,;№N ]+\) |\(см\. текст [а-яА-Я0-9.,;№N ]+\) |\(абзац введен [а-яА-Я0-9.,;№N ]+\) )', r'', result)
# Удаление номеров ссылок и сносок и их текста
result = re.sub(r'(?<=[а-яА-Я])([0-9*]+)', r'', result)
result = re.sub(r'([0-9*] *Далее [а-яА-Я0-9-–,:;№N ]+\. )', r'', result)
# Удаление символа дефиса при переносе слова на следующую строку
result = re.sub(r'(\xad) ', r'', result)
# Удаление фразы
result = re.sub(r'Список изменяющих документов', r'', result)
# Добавление символа новой строки на разделы
result = re.sub(r' (I+V*\.) ', r'\n\1', result)
# Удаление пробела после дефиса
result = re.sub(r'(\w-) ', r'\1', result)
# print(result)

# _______________________________________________________________________________________________________________________
# Преобразование строки в список
list_results = result.split('\n')

# Сохранение в словарь по разделам (I, II, III...)
moduls_doc = {"Титульная информация": list_results[0],
              "Общие положения": list_results[1],
              "Задачи": list_results[2],
              "Функции": list_results[3],
              "Руководство": list_results[4]}

for key in moduls_doc:
    # Добавление символа новой строки по нумерации пунктов
    moduls_doc[key] = re.sub(r'(?<=[а-яА-Я.,:;)]) (?=([0-9]{1,2}[.)])([0-9]{1,2}\))?[^(\d{2}.\d{2}.\d{4})])', r'\n', moduls_doc[key])
    # Добавление символа новой строки после : и ;
    moduls_doc[key] = re.sub(r'(?<=[:;]) (?=[а-яА-Я])', r'\n-', moduls_doc[key])
    # Добавление символа новой строки после точки
    moduls_doc[key] = re.sub(r'([а-я]{2,}\)?\.)( )(?=[А-Я])', r'\1\n', moduls_doc[key])
    # Добавление символа новой строки после слова
    moduls_doc[key] = re.sub(r'(Утверждено) ', r'\1\n', moduls_doc[key])
# print('-' * 80)
# for key in moduls_doc:
#     print(f'modul |{key: <20}| from moduls_docs => {moduls_doc[key]}')
#     print('-' * 80)
# _______________________________________________________________________________________________________________________

# Определение ОВУ
ovy_name = re.search(r'(?<=Об\sутверждении\s[Пп]оложения\sо\s)[а-яА-Я ]*(?=\sГенерального\sштаба\sВооруженных\sСил\sРоссийской\sФедерации)', moduls_doc['Титульная информация'])
# Преобразование слова в родительный падеж
morph = pymorphy2.MorphAnalyzer()
ovy_name = ovy_name[0].split(' ')
# Переменная, в которой сохраняются результаты пробразования
ovy_name_gent = ''
#В цикле проходим по каждому слову
for word in ovy_name:
    # Обработка слова при помощи pymorphy2
    refactor_word = morph.parse(word)[0]
    # Преобразование слова в родительный падеж
    case = refactor_word.inflect({'gent'})
    # Если ovy_name_gent пустая переменная - добавить слово с заглавной буквой,
    # в остальных случаях конкатенация с пробелом
    if ovy_name_gent:
        ovy_name_gent += ' ' + case.word
    else:
        ovy_name_gent += case.word.title()
ovy_name = ovy_name_gent
# _______________________________________________________________________________________________________________________

# Удаление заглавных строк разделов
moduls_doc['Задачи'] = moduls_doc['Задачи'].split('\n')[2:]
moduls_doc['Задачи'] = '\n'.join(moduls_doc['Задачи'])
moduls_doc['Функции'] = moduls_doc['Функции'].split('\n')[2:]
moduls_doc['Функции'] = '\n'.join(moduls_doc['Функции'])

# Дополнение нумерации задач
Duties_z = []
count_z = 1
for match in re.finditer(r'(?P<Duti_z>[0-9]+\)\s(?P<Duti_z_not_number>.*[;.]))', moduls_doc['Задачи']):
    if match["Duti_z"]:
        duti_z = ['Задача', count_z, ovy_name, match["Duti_z_not_number"]]
        Duties_z.append(f"{duti_z[0]} {duti_z[1]} {duti_z[2]}| {duti_z[3]}")
        count_z += 1
Duties_z = '\n'.join(Duties_z)
moduls_doc['Задачи'] = Duties_z

# Задание и дополнение нумерации функций
Duties = []
count_f = 1
count_ppf = 1
duti_f = []
duti_pf = []
duti_ppf = []
for match in re.finditer(r'(?P<Duti_f>[0-9]+\)\s(?P<Duti_f_not_number>.*:))|(?P<Duti_pf>-.*:)+|(?P<Duti_ppf>-?.*[;.])+', moduls_doc['Функции']):
    global count_pf
    if match["Duti_f"]:
        duti_f = ['Функция', count_f, ovy_name, match["Duti_f_not_number"]]
        Duties.append(f"{duti_f[0]} {duti_f[1]} {duti_f[2]}| {duti_f[3]}")
        count_f += 1
        count_pf = 1
        count_ppf = 1
    elif match["Duti_pf"]:
        duti_pf = [count_f - 1, count_pf, match["Duti_pf"][1:]]
        Duties.append(f'{duti_pf[0]}_{duti_pf[1]}| {duti_pf[2]}')
        count_pf += 1
        count_ppf = 1
    elif match["Duti_ppf"]:
        if duti_pf:
            if duti_pf[0] == count_f - 1:
                duti_ppf = ['Функция', count_f - 1, count_pf, count_ppf, ovy_name, match["Duti_ppf"][1:]]
                Duties.append(f'{duti_ppf[0]} {count_f - 1}_{duti_pf[1]}_{count_ppf} {duti_ppf[4]}| {duti_pf[2][:-1]}| {duti_ppf[5]}')
                count_ppf += 1
            else:
                duti_ppf = ['Функция', count_f - 1, count_ppf, ovy_name, match["Duti_ppf"][1:]]
                Duties.append(f'{duti_ppf[0]} {count_f - 1}_{count_ppf} {duti_ppf[3]} | {duti_ppf[4]}')
                count_ppf += 1
        else:
            duti_ppf = ['Функция', count_f - 1, count_ppf, ovy_name, match["Duti_ppf"][1:]]
            Duties.append(f'{duti_ppf[0]} {count_f - 1}_{count_ppf} {duti_ppf[3]}| {duti_ppf[4]}')
            count_ppf += 1
Duties = '\n'.join(Duties)
moduls_doc['Функции'] = Duties
# print(Duties)

print('-' * 80)
for key in moduls_doc:
    if key != "Руководство":
        print(moduls_doc[key])
        print('-' * 80)
# _______________________________________________________________________________________________________________________

#Сохранение результатов в формате docx
document = Document()

# Добавление заголовка. Уровень от 0 до 9 (по умолчанию, 1)
document.add_heading('Результат работы алгоритма', 0)
document.add_heading(pdf_file[1], 1)

# Добавление заголовка "Задачи"и соответствующих параграфов
document.add_heading(f'Задачи {ovy_name}', level=2)
moduls_doc['Задачи'] = moduls_doc['Задачи'].split('\n')
for paragraph in moduls_doc['Задачи']:
    document.add_paragraph(paragraph)

# Добавление заголовка "Функции" и соответствующих параграфов
document.add_heading(f'Функции {ovy_name}', level=2)
moduls_doc['Функции'] = moduls_doc['Функции'].split('\n')
for paragraph in moduls_doc['Функции']:
    document.add_paragraph(paragraph)


document.save('Форма представления распознанных обязанностей.docx')
print('Done')