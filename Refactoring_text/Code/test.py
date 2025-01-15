import fitz
import re
from docx import Document
import pymorphy2


# Цвета для вывода
class bcolors:
    OKGREEN = '\033[92m'
    YELLOW = '\033[93m'
    ENDC = '\033[0m'
    UND = '\033[4m'
    BOLD = '\033[1m'
    CYAN = '\033[96m' # {bcolors.CYAN}{....}{bcolors.ENDC}


def title_list():
    pass


def general_provisions():
    pass


# Дополнение нумерации полномочий
def authorities(text_from_key_a):
    # Корректировка полномочия 33.9
    text_from_key_a = re.sub('\n8\.\sМинобороны России[-а-яА-Я0-9().,:; \n]+', '', text_from_key_a)
    # Удаление лишних строк после полномочий
    text_from_key_a = re.sub('(?<=или иной государственной службе;)\n-(?=утверждает форму одежды обучающихся)',' ', text_from_key_a)
    # Итоговый список полномочий
    Duties_a = []
    # Счетчик номера выявленного текста полномочия
    count_ppa = 1
    # Счётчик выявленния паттернов
    count_pattern = 0
    # Создание списов для хранения результатов последнего распознавания по шаблону
    # Список с полномочием
    duti_a = []
    # Список с заголовком полномочия с пунктами
    duti_pa = []
    # Список с пунктом полномочия
    duti_ppa = []
    for match in re.finditer(r'(?P<Duti_a>(?P<Duti_a_number>[0-9]+(\.[0-9]+)?)\)\s(?P<Duti_a_not_number>.*[;.]))|'
                             r'(?P<Duti_pa>(?P<Duti_pa_number>[0-9]+(\.[0-9]+)?)\)\s(?P<Duti_pa_not_number>.*:))|'
                             r'(?P<Duti_ppa>-?.*[;.])', text_from_key_a):
        # Поиск полномочий
        if match["Duti_a"]:
            count_pattern += 1
            # "Полномочие", номер полномочия, ОВУ, текст полномочия, номер выявленного паттерна, уровень полномочия
            duti_a = ['Полномочие', match["Duti_a_number"], ovy_name, match["Duti_a_not_number"][:-1] + '.']
            Duties_a.append([duti_a[0], duti_a[1], duti_a[2], duti_a[3]])
            # print(f'{duti_a[0]} {duti_a[1]} {duti_a[2]} - {duti_a[3]}')
            count_ppa = 1
        # Поиск под полномочий
        elif match["Duti_pa"]:
            count_pattern += 1
            # Номер полномочия, номер под полномочия, текст под полномочия, уровень под полномочия
            duti_pa = ['Полномочие', match["Duti_pa_number"], ovy_name, match["Duti_pa_not_number"]]
            # print(f'{duti_pa[0]} {duti_pa[1]} {duti_pa[2]} - {duti_pa[3]}')
            count_ppa = 1
        # Поиск текста под полномочий
        elif match["Duti_ppa"]:
            count_pattern += 1
            # Номер полномочия, номер под полномочия, текст под полномочия, уровень под полномочия
            duti_ppa = ['Полномочие', duti_pa[1], count_ppa, ovy_name, duti_pa[3], match["Duti_ppa"][1:-1] + '.']
            Duties_a.append([duti_ppa[0], duti_ppa[1], duti_ppa[2], duti_ppa[3], duti_ppa[4], duti_ppa[5]])
            # print(f'\t{duti_ppa[1]}_{duti_ppa[2]} - {duti_ppa[5]}')
            count_ppa += 1
    print(f'count_pattern - <{count_pattern}>')
    return Duties_a


# Дополнение нумерации задач
def tasks(text_from_key_z: str):
    Duties_z = []
    count_z = 1
    for match in re.finditer(r'(?P<Duti_zn>[0-9]+\)\s(?P<Duti_zn_not_number>.*[;.]))|'
                             r'(?P<Duti_z>(?P<Duti_z_not_number>.*[;.]))', text_from_key_z):
        if match["Duti_zn"]:
            Duties_z.append(['Задача', count_z, ovy_name, match["Duti_zn_not_number"]])
            count_z += 1
        elif match["Duti_z"]:
            Duties_z.append(['Задача', count_z, ovy_name, match["Duti_z_not_number"][1:]])
            count_z += 1
    return Duties_z


def tasks_MO(text_from_key_tMO: str):
    text_from_key_tMO = re.search(r'((?<=:)[-а-яА-Я0-9().,:; \n]+(?=4\.))', text_from_key_tMO)
    Duties_tMO = []
    for match in re.finditer(r'(?P<Duti_tMO>(?P<Duti_tMO_number>[0-9]+(\.[0-9]+)?)\)\s(?P<Duti_tMO_not_number>.*[;.]))', text_from_key_tMO[0]):
        Duties_tMO.append(['Задача', match["Duti_tMO_number"], ovy_name, match["Duti_tMO_not_number"]])
        # print(f'{Duties_tMO[-1][0]} {Duties_tMO[-1][1]} {Duties_tMO[-1][2]} - {Duties_tMO[-1][3]}')
    return Duties_tMO

# Дополнение нумерации функций
def functions(text_from_key_f: str):
    # Итоговый список функций
    Duties_f = []
    # Счетчик номера функции
    count_f = 1
    # Подсчёт добавления функции в "Duties_f", кроме одной и той же подфункции
    count_uf = 1
    # Счетчик номера подфункции
    count_pf = 0
    # Счетчик номера выявленного текста функции
    count_ppf = 1
    # Счетчик уровня функции
    count_deep = 1
    # Счётчик выявленния паттернов
    count_pattern = 0
    # Создание списов для хранения результатов последнего распознавания по шаблону
    # Список с шапкой функции
    duti_f = []
    # Список с подфункцией без буквы в нумерации
    duti_pf = []
    # Список с подфункцией с буквой в нумерации
    duti_pfl = []
    # Список с текстом функции без буквы в нумерации
    duti_ppf = []
    # Список с текстом функции с буквой в нумерации
    duti_ppfl = []
    for match in re.finditer(r'(?P<Duti_f>[0-9]+\)\s(?P<Duti_f_not_number>.*:))|'
                             r'(?P<Duti_pfl>(?P<Duti_pf_letter>[а-я])\).*:)+|'
                             r'(?P<Duti_pf>-?.*:)+|'
                             r'(?P<Duti_ppfl>(?P<Duti_ppf_letter>[а-я])\).*[;.])+|'
                             r'(?P<Duti_ppf>-?.*[;.])+', text_from_key_f):
        # Поиск функций
        if match["Duti_f"]:
            # Обнуляется уровень функции
            count_pattern += 1
            count_deep = 1
            # "Функция", номер функции, ОВУ, текст функции, номер выявленного паттерна, уровень функции
            duti_f = ['Функция', count_f, ovy_name, match["Duti_f_not_number"][:-1], count_pattern, count_deep]
            # Все print() в теле цикла позволяют вывести функции по уровням
            print(f'{duti_f[0]} {duti_f[1]} {duti_f[2]} {duti_f[3]}')
            print(f'count_deep - {count_deep}')
            count_f += 1
            count_uf = 1
            count_pf = 0
            count_ppf = 1
        # Поиск подфункций
        elif match["Duti_pf"] or match["Duti_pfl"]:
            # Если номер подфункции с буквой
            if match["Duti_pfl"]:
                # Если последняя добавленная в Duties_f функция с буквой
                if Duties_f and type(match["Duti_pf_letter"]) == type(Duties_f[-1][3]):
                    # Счетчик уровня функции приравнивается к
                    # счетчику уровня последней добавленной в Duties_f функции
                    count_deep = Duties_f[-1][-1]
                    count_pattern += 1
                    # Номер функции, номер подфункции (согласно номеру последней добавленной в Duties_f функции),
                    # буква из документа, текст подфункции, уровень подфункции
                    duti_pfl = [count_f - 1, Duties_f[-1][2], match["Duti_pf_letter"], match["Duti_pfl"][3:-1],
                                count_pattern, count_deep]
                    print("\t" * (count_deep - 1) + f'{duti_pfl[0]}_{duti_pfl[1]}_{duti_pfl[2]}| {duti_pfl[3]}')
                    print("\t" * (count_deep - 1) + f'count_deep - {count_deep}')
                # В остальных случаях
                else:
                    # Уровень увеличивается на 1
                    count_deep = duti_f[-1] + 1
                    count_pattern += 1
                    # Номер функции, буква из документа, текст подфункции, уровень подфункции
                    duti_pfl = [count_f - 1, match["Duti_pf_letter"], match["Duti_pfl"][3:], count_pattern, count_deep]
                    print(f'\t{duti_pfl[0]}_{duti_pfl[1]}| {duti_pfl[2]}')
                    print(f'\tcount_deep - {count_deep}')
            # Если номер подфункции без буквы
            else:
                # Уровень увеличивается на 1
                count_pattern += 1
                count_deep = duti_f[-1] + 1
                # Номер функции, номер подфункции, текст подфункции, уровень подфункции
                duti_pf = [count_f - 1, count_uf, match["Duti_pf"][1:-1], count_pattern, count_deep]
                print(f'\t{duti_pf[0]}_{duti_pf[1]}| {duti_pf[2]}')
                print(f'\tcount_deep - {count_deep}')
                count_uf += 1
            count_pf += 1
            count_ppf = 1
        # Поиск текста функций
        elif match["Duti_ppf"] or match["Duti_ppfl"]:
            # Если последний распознанный паттерн - функция
            if duti_f[-2] == count_pattern:
                count_deep = 2
                count_pattern += 1
                if match["Duti_ppf"]:
                    # "Функция", номер функции, номер текста функции, ОВУ, функция, текст функции, номер распознанного паттерна, уровень функции
                    duti_ppf = ['Функция', count_f - 1, count_uf, ovy_name, duti_f[3], match["Duti_ppf"][1:], count_pattern, count_deep]
                    Duties_f.append([duti_ppf[0], duti_ppf[1], duti_ppf[2], duti_ppf[3], duti_ppf[4], duti_ppf[5], duti_ppf[6], duti_ppf[7]])
                    print("\t" * (count_deep - 1) + f'{duti_ppf[1]}.{duti_ppf[2]}) {duti_ppf[-3]}')
                    count_uf += 1
                elif match["Duti_ppfl"]:
                    # "Функция", номер функции, буква из документа, ОВУ, функция, подфункция, текст функции, номер распознанного паттерна, уровень функции
                    duti_ppfl = ['Функция', count_f - 1, match['Duti_ppf_letter'], ovy_name, duti_f[3], match["Duti_ppfl"][3:], count_pattern, count_deep]
                    Duties_f.append([duti_ppfl[0], duti_ppfl[1], duti_ppfl[2], duti_ppfl[3], duti_ppfl[4], duti_ppfl[5], duti_ppfl[6], duti_ppfl[7]])
                    print("\t" * (count_deep - 1) + f'{duti_ppfl[1]}.{duti_ppfl[2]}) {duti_ppfl[-3]}')
                print("\t" * (count_deep - 1) + f'count_deep - {count_deep}')
            # Если последний распознанный паттерн - подфункция без буквы в номере
            elif duti_pf and duti_pf[-2] == count_pattern:
                count_deep = duti_pf[-1] + 1
                count_pattern += 1
                if match["Duti_ppf"]:
                    # "Функция", номер функции, номер подфункции, номер текста функции, ОВУ, функция, подфункция, текст функции, номер распознанного паттерна, уровень функции
                    duti_ppf = ['Функция', count_f - 1, count_uf - 1, count_ppf, ovy_name, duti_f[3], duti_pf[2], match["Duti_ppf"][1:], count_pattern, count_deep]
                    Duties_f.append([duti_ppf[0], duti_ppf[1], duti_ppf[2], duti_ppf[3], duti_ppf[4], duti_ppf[5], duti_ppf[6], duti_ppf[7], duti_ppf[8], duti_ppf[9]])
                    print("\t" * (count_deep - 1) + f'{duti_ppf[1]}.{duti_ppf[2]}.{duti_ppf[3]}) {duti_ppf[-3]}')
                    count_ppf += 1
                elif match["Duti_ppfl"]:
                    # "Функция", номер функции, номер подфункции, буква из документа, ОВУ, функция, подфункция, текст функции, номер распознанного паттерна, уровень функции
                    duti_ppfl = ['Функция', count_f - 1, count_uf - 1, match['Duti_ppf_letter'], ovy_name, duti_f[3], duti_pf[2], match["Duti_ppfl"][3:], count_pattern, count_deep]
                    Duties_f.append([duti_ppfl[0], duti_ppfl[1], duti_ppfl[2], duti_ppfl[3], duti_ppfl[4], duti_ppfl[5], duti_ppfl[6], duti_ppfl[7], duti_ppfl[8], duti_ppfl[9]])
                    print("\t" * (count_deep - 1) + f'{duti_ppfl[1]}.{duti_ppfl[2]}.{duti_ppfl[3]}) {duti_ppfl[-3]}')
                print("\t" * (count_deep - 1) + f'count_deep - {count_deep}')
            # Если последний распознанный паттерн - подфункция с буквой в номере
            elif duti_pfl and duti_pfl[-2] == count_pattern:
                count_deep = duti_pfl[-1] + 1
                count_pattern += 1
                if match["Duti_ppf"]:
                    if count_deep == 4:
                        # "Функция", номер функции, номер подфункции, буква из документа, номер текста функции, ОВУ, функция, подфункция, текст функции, номер распознанного паттерна, уровень функции
                        duti_ppf = ['Функция', count_f - 1, duti_pf[1], duti_pfl[2], count_ppf, ovy_name, duti_f[3], duti_pf[2], duti_pfl[2], match["Duti_ppf"][1:], count_pattern, count_deep]
                        Duties_f.append(
                            [duti_ppf[0], duti_ppf[1], duti_ppf[2], duti_ppf[3], duti_ppf[4], duti_ppf[5], duti_ppf[6], duti_ppf[7], duti_ppf[8], duti_ppf[9], duti_ppf[10], duti_ppf[11]])
                        print("\t" * (count_deep - 1) + f'{duti_ppf[1]}.{duti_ppf[2]}.{duti_ppf[3]}.{duti_ppf[4]}) {duti_ppf[-3]}')
                        count_ppf += 1
                    elif count_deep == 3:
                        # "Функция", номер функции, буква из документа, номер текста функции, ОВУ, функция, подфункция, текст функции, номер распознанного паттерна, уровень функции
                        duti_ppf = ['Функция', count_f - 1, duti_pfl[1], count_ppf, ovy_name, duti_f[3], duti_pfl[2], match["Duti_ppf"][1:], count_pattern, count_deep]
                        Duties_f.append([duti_ppf[0], duti_ppf[1], duti_ppf[2], duti_ppf[3], duti_ppf[4], duti_ppf[5], duti_ppf[6], duti_ppf[7], duti_ppf[8], duti_ppf[9]])
                        print("\t" * (count_deep - 1) + f'{duti_ppf[1]}.{duti_ppf[2]}.{duti_ppf[3]}) {duti_ppf[-3]}')
                        count_ppf += 1
                elif match["Duti_ppfl"]:
                    pass
                print("\t" * (count_deep - 1) + f'count_deep - {count_deep}')
            # Если последний распознанный паттерн - текст функции без буквы в номере
            elif duti_ppf and duti_ppf[-2] == count_pattern:
                count_pattern += 1
                if match["Duti_ppf"]:
                    count_deep = duti_ppf[-1]
                    if count_deep == 4:
                        # "Функция", номер функции, номер подфункции, буква из документа, номер текста функции, ОВУ, функция, подфункция1, подфункция2, текст функции, номер распознанного паттерна, уровень функции
                        duti_ppf = ['Функция', count_f - 1, duti_pf[1], duti_pfl[2], count_ppf, ovy_name, duti_f[3], duti_pf[2], duti_pfl[2], match["Duti_ppf"][1:], count_pattern, count_deep]
                        Duties_f.append([duti_ppf[0], duti_ppf[1], duti_ppf[2], duti_ppf[3], duti_ppf[4], duti_ppf[5], duti_ppf[6], duti_ppf[7], duti_ppf[8], duti_ppf[9], duti_ppf[10], duti_ppf[11]])
                        print("\t" * (count_deep - 1) + f'{duti_ppf[1]}.{duti_ppf[2]}.{duti_ppf[3]}.{duti_ppf[4]}) {duti_ppf[-3]}')
                    elif count_deep == 3:
                        if duti_pfl:
                            # "Функция", номер функции, номер подфункции, номер текста функции, ОВУ, функция, подфункция, текст функции, номер распознанного паттерна, уровень функции
                            duti_ppf = ['Функция', count_f - 1, duti_pfl[1], count_ppf, ovy_name, duti_f[3], duti_pfl[2], match["Duti_ppf"][1:], count_pattern, count_deep]
                        if duti_pf:
                            # "Функция", номер функции, номер подфункции, номер текста функции, ОВУ, функция, подфункция, текст функции, номер распознанного паттерна, уровень функции
                            duti_ppf = ['Функция', count_f - 1, count_uf - 1, count_ppf, ovy_name, duti_f[3], duti_pf[2], match["Duti_ppf"][1:], count_pattern, count_deep]
                        Duties_f.append([duti_ppf[0], duti_ppf[1], duti_ppf[2], duti_ppf[3], duti_ppf[4], duti_ppf[5], duti_ppf[6], duti_ppf[7], duti_ppf[8], duti_ppf[9]])
                        print("\t" * (count_deep - 1) + f'{duti_ppf[1]}.{duti_ppf[2]}.{duti_ppf[3]}) {duti_ppf[-3]}')
                    elif count_deep == 2:
                        # "Функция", номер функции, номер текста функции, ОВУ, функция, текст функции, номер распознанного паттерна, уровень функции
                        duti_ppf = ['Функция', count_f - 1, count_uf, ovy_name, duti_f[3], match["Duti_ppf"][1:], count_pattern, count_deep]
                        Duties_f.append([duti_ppf[0], duti_ppf[1], duti_ppf[2], duti_ppf[3], duti_ppf[4], duti_ppf[5], duti_ppf[6], duti_ppf[7]])
                        print("\t" * (count_deep - 1) + f'{duti_ppf[1]}.{duti_ppf[2]}) {duti_ppf[-3]}')
                        count_uf += 1
                    count_ppf += 1
                elif match["Duti_ppfl"]:
                    count_deep = duti_pfl[-1]
                    if count_deep == 3:
                        # "Функция", номер функции, номер подфункции, буква из документа, ОВУ, заголовок функции, текст подфункции, текст функции, номер распознанного паттерна, уровень функции
                        duti_ppfl = ['Функция', count_f - 1, count_uf - 1, match['Duti_ppf_letter'], ovy_name, duti_f[3], duti_pf[2], match["Duti_ppfl"][3:], count_pattern, count_deep]
                        Duties_f.append([duti_ppfl[0], duti_ppfl[1], duti_ppfl[2], duti_ppfl[3], duti_ppfl[4], duti_ppfl[5], duti_ppfl[6], duti_ppfl[7], duti_ppfl[8], duti_ppfl[9]])
                    elif count_deep == 2:
                        # "Функция", номер функции, буква из документа, ОВУ, заголовок функции, текст функции, номер распознанного паттерна, уровень функции
                        duti_ppfl = ['Функция', count_f - 1, match['Duti_ppf_letter'], ovy_name, duti_f[3], match["Duti_ppfl"][3:], count_pattern, count_deep]
                        Duties_f.append([duti_ppfl[0], duti_ppfl[1], duti_ppfl[2], duti_ppfl[3], duti_ppfl[4], duti_ppfl[5], duti_ppfl[6], duti_ppfl[7]])
                    print("\t" * (count_deep - 1) + f'{duti_ppfl[1]}.{duti_ppfl[2]}) {duti_ppfl[-3]}')
                print("\t" * (count_deep - 1) + f'count_deep - {count_deep}')
            # Если последний распознанный паттерн - текст функции с буквой в номере
            elif duti_ppfl and duti_ppfl[-2] == count_pattern:
                count_pattern += 1
                if match["Duti_ppf"]:
                    if duti_pf:
                        count_deep = duti_pf[-1]
                    elif duti_pfl:
                        count_deep = duti_pfl[-1] - 1
                    # "Функция", номер функции, номер текста функции, ОВУ, функция, текст функции, номер распознанного паттерна, уровень функции
                    duti_ppf = ['Функция', count_f - 1, count_uf, ovy_name, duti_f[3], match["Duti_ppf"][1:], count_pattern, count_deep]
                    Duties_f.append([duti_ppf[0], duti_ppf[1], duti_ppf[2], duti_ppf[3], duti_ppf[4], duti_ppf[5], duti_ppf[6], duti_ppf[7]])
                    print("\t" * (count_deep - 1) + f'{duti_ppf[1]}.{duti_ppf[2]}) {duti_ppf[-3]}')
                    count_uf += 1
                elif match["Duti_ppfl"]:
                    count_deep = duti_ppfl[-1]
                    if count_deep == 3:
                        # "Функция", номер функции, номер подфункции, буква из документа, ОВУ, заголовок функции, текст подфункции, текст функции, номер распознанного паттерна, уровень функции
                        duti_ppfl = ['Функция', count_f - 1, count_uf - 1, match['Duti_ppf_letter'], ovy_name, duti_f[3], duti_pf[2], match["Duti_ppfl"][3:], count_pattern, count_deep]
                        Duties_f.append([duti_ppfl[0], duti_ppfl[1], duti_ppfl[2], duti_ppfl[3], duti_ppfl[4], duti_ppfl[5], duti_ppfl[6], duti_ppfl[7], duti_ppfl[8], duti_ppfl[9]])
                        print("\t" * (count_deep - 1) + f'{duti_ppfl[1]}.{duti_ppfl[2]}.{duti_ppfl[3]}) {duti_ppfl[-3]}')
                    elif count_deep == 2:
                        # "Функция", номер функции, буква из документа, ОВУ, заголовок функции, текст функции, номер распознанного паттерна, уровень функции
                        duti_ppfl = ['Функция', count_f - 1, match['Duti_ppf_letter'], ovy_name, duti_f[3], match["Duti_ppfl"][3:], count_pattern, count_deep]
                        Duties_f.append([duti_ppfl[0], duti_ppfl[1], duti_ppfl[2], duti_ppfl[3], duti_ppfl[4], duti_ppfl[5], duti_ppfl[6], duti_ppfl[7]])
                        print("\t" * (count_deep - 1) + f'{duti_ppfl[1]}.{duti_ppfl[2]}) {duti_ppfl[-3]}')
                print("\t" * (count_deep - 1) + f'count_deep - {count_deep}')
    print(f'Было выявлено {count_pattern} шаблонов.')
    return Duties_f


pdf_file = ["Положение о Минобороны.pdf",
            "ПМО 2014 г. № 400дсп Положение о ГОУ.pdf",
            "ПМО 2016 г. № 260дсп Положение о ГК СВ.pdf",
            "Приказ Министр обороны Российской Федерации от 16.10.2018 № 580 (в редакции от 27.01.2024 № 41).pdf",
            "ПМО 2017 г. № 600 Положение о ВТУ.pdf",
            "Приказ Министр обороны Российской Федерации от 16.01.2016 № 4 Положение о ДФМ ГОЗ оригинал.pdf",
            "ПМО 2016 г. № 240 Положение о ГОМУ.pdf",
            "Приказ Министр обороны Российской Федерации от 03.09.2020 № 444 Положение о ДВИ оригинал.pdf",
            "Приказ Министр обороны Российской Федерации от 06.12.2021 № 727 Положение о ДТО.pdf",
            "Приказ Министр обороны Российской Федерации от 15.11.2018 № 575 Положение о ПД.pdf"]

pdf_path = "docs/" + pdf_file[0]
text = ''
file = fitz.open(pdf_path)
for pageNum, page in enumerate(file.pages()):
    text += page.get_text()
# Создаётся строковая переменная из всех данных
result = text.replace('\n', ' ')
# print(result)
# print('-' * 80)
# Удаление номеров страниц (Дублируется чтобы убрать непонятно откуда взявшиеся лишние номера)
result = re.sub(r'((?<![«№N-]) +[0-9]{1,2} +(?!янв|фев|мар|апр|май|июн|июл|авг|сент|окт|ноя|дек|прил|(к )|»))', ' ', result)
result = re.sub(r'((?<![«№N-]) +[0-9]{1,2} +(?!янв|фев|мар|апр|май|июн|июл|авг|сент|окт|ноя|дек|прил|(к )))', ' ', result)
# Удаление лишних пробелов
result = re.sub(r' {2,}', ' ', result)
result = re.sub(r'(?<=[а-яА-Я]) (?=;)', '', result)
result = re.sub(r'(?<=[а-я]) (?=\))', '', result)
# Удаление нижних подчеркиваний
result = re.sub(r'_+', '', result)
# Удаление символа дефиса при переносе слова на следующую строку
# result = re.sub(r'(\xad) ', r'', result)
# print(result)
# print('-' * 80)
# Удаление текста в скобках о редакциях, пунктах и т.д.
result = re.sub(r'(\((п\. )?в ред\. [а-яА-Я0-9.,;№N ]+\) |'
                r'\(п\. [0-9]*\.?[0-9]* [а-яА-Я0-9.,;№N ]+\) |'
                r'\(пп\. [0-9]*\.?[0-9]* [а-яА-Я0-9.,;№N ]+\) |'
                r'\(пп\. в ред\. [а-яА-Я0-9.,;№N ]+\) |'
                r'\(см\. текст [а-яА-Я0-9.,;№N ]+\) |'
                r'\(абзац введен [а-яА-Я0-9.,;№N ]+\) )', r'', result)
# print(result)
# print('-' * 80)
# Удаление символа дефиса при переносе слова на следующую строку
result = re.sub(r'(\xad) (?!мобилизац)', r'', result)
result = re.sub(r'(\xad) (?=мобилизац)', r'\1', result)
# Удаление мусорных фраз
result = re.sub(r'НОРМАТИВНЫЙ ', r'', result)
result = re.sub(r'Список изменяющих документов ', r'', result)
result = re.sub(r'Документ\sпредоставлен\sКонсультантПлюс(\s(www\.consultant\.ru\s)*Дата\sсохранения:\s\d{1,2}\.\d{1,2}\.\d{4})? ', '', result)
result = re.sub(r'КонсультантПлюс', '', result)
result = re.sub(r'надежная\sправовая\sподдержка', '', result)
result = re.sub(r'www\.consultant\.ru', '', result)
# Удаление пробела после дефиса
result = re.sub(r'(\w-) ', r'\1', result)
# print(result)
# print('-' * 80)
# Удаление номеров ссылок и сносок и их текста
result = re.sub(r'(?<=[а-яА-Я)])([*]+)\s?', r'', result)
result = re.sub(r'([0-9*-] *Далее [а-яА-Я0-9-–,:;№N ]+\. )', r'', result)
result = re.sub(r'([*“-]{0,3}\s?(Далее в тексте|Здесь и далее|Система транспортного обеспечения|'
                r'В части номенклатуры|Порядок взаимодействия|Во взаимодействии|'
                r'Под эксплуатационными|Под военно-автомобильными дорогами)[^0-9]+?\. (?=[а-я]))',
                r'', result, re.MULTILINE)
result = re.sub(r"' Здесь и далее в тексте настоящего Положения под воинскими автомобильными перевозками понимаются воинские перевозки, выполняемые автомобильным транспортом общего пользования \(за исключением такси\)\. |"
                r"“ Порядок взаимодействия Министерства обороны с другими войсками, воинскими формированиями и органами по вопросам транспортного обеспечения регулируется нормативными правовыми актами Российской Федерации и приказами Министра обороны Российской Федерации\. |"
                r"\* Система транспортного обеспечения - совокупность взаимоувязанных и согласованных по задачам и целям органов военного управления транспортного обеспечения, военно-транспортных органов, автомобильных и дорожных войск, соединений и организаций вспомогательного флота, предназначенных для осуществления транспортных потребностей Вооруженных Сил с использованием железнодорожного, воздушного, морского, внутреннего водного и автомобильного транспорта\. |"
                r"Здесь и далее в тексте настоящего Положения под перевозками автомобильным транспортом Вооруженных Сил понимаются мобилизационные, оперативные, снабженческие и эвакуационные перевозки, выполняемые автомобильными войсками\. |"
                r"В части номенклатуры «морские и рейдовые суда обеспечения» - совместно с ГК ВМФ\. |"
                r"['*] Во взаимодействии с ГК ВМФ\. |"
                r"\* Под эксплуатационными специальными формированиями понимаются военно-эксплуатационные отделения, колонны тепловозов особого резерва\. |"
                r"\* Под военно-автомобильными дорогами Центра понимаются военноавтомобильные дороги, связывающие важные экономические регионы страны, центры материально-технического обеспечения, выгрузочные станции и аэродромы с районами размещения стационарной инфраструктуры материально-технического обеспечения и военно-автомобильными дорогами военных округов \(флотов\)\. |"
                r"\* Правовая экспертиза проектов документов, указанных в подпунктах «а», «б», «в» подпункта 1, абзацах четвертом подпункта 3, втором » третьем подпункта 5, четвертом. И пятом подпункта б, десятом подпункта пункта настоящего Положения, проводится в установленном порядке, после их согласования со всеми заинтересованными органами военного управления\. 5 ", r'', result)
# Частные случаи для Положения о ДТО №727
result = re.sub(r"(?<=а\) сбор, обобщение и анализ обеспеченности Вооруженных Сил материальными ценностями, а также потребностей в работах, товарах, услугах по закрепленной номенклатуре, подготовка предложений для принятия соответствующих решений; )"
                r"6(?=\) подготовка в пределах своей компетенции предложений в планы и программы обеспечения Вооруженных Сил материальными ценностями)", 'б', result)
result = re.sub(r'(?<=; ).(?=и\))', r'', result)
# Частные случаи для Положения о ПД №575
result = re.sub(r"(?<=III),(?= Функции)", ".", result)
result = re.sub(r'(?<=Функции Правового департамента 8),(?= Правовой департамент осуществляет следующие функции:)', r'.', result)
result = re.sub(r"(?<=правовое и организационное обеспечение взаимодействия заместителей Министра обороны Российской Федерации, руководства центральных):(?= органов военного управления с полномочными представителями Президента Российской Федерации и Правительства Российской Федерации)", "", result)
# Добавление символа новой строки на разделы
result = re.sub(r' (I+V*\.) ', r'\n\1', result)
# Добавление символа новой строки после слова
result = re.sub(r'(Утверждено|г. Москва генерал армии С\.Шойгу) ', r'\1\n', result)
# _______________________________________________________________________________________________________________________

list_results = result.split('\n')
for listsr in list_results:
    print(listsr)

# Сохранение в словарь по разделам (I, II, III...)
moduls_doc = {}
for list_result in list_results:
    if re.search(r'I\.\s?Общие положения', list_result):
        moduls_doc['Общие положения'] = list_result
    elif re.search(r'II\.\s?Полномочия', list_result):
        moduls_doc['Полномочия'] = list_result
    elif re.search(r'II\.\s?Основные задачи', list_result):
        moduls_doc['Задачи'] = list_result
    elif re.search(r'III\.\s?Функции', list_result):
        moduls_doc['Функции'] = list_result
    elif re.search(r'IV\.\s?Руководство', list_result):
        moduls_doc['Руководство'] = list_result
    elif re.search(r'III[.,]\s?Организация деятельности', list_result):
        moduls_doc['Организация деятельности'] = list_result
    elif re.search(r'УКАЗ ПРЕЗИДЕНТА', list_result):
        moduls_doc['Указ Президента'] = list_result
    else:
        moduls_doc['Титульная информация'] = list_result

for key in moduls_doc:
    # Добавление символа новой строки по нумерации пунктов
    moduls_doc[key] = re.sub(r'(?<=[а-яА-Я.,:;)]) (?=([0-9]{1,2}[.)])([0-9]{1,2}\))?[^(\d{2}.\d{2}.\d{4})])', r'\n', moduls_doc[key])
    # moduls_doc[key] = re.sub(r'(?<=ВОПРОСЫ МИНИСТЕРСТВА ОБОРОНЫ РОССИЙСКОЙ ФЕДЕРАЦИИ)  (?=1\.)', r'\n', moduls_doc[key])
    # Добавление символа новой строки после : и ;
    moduls_doc[key] = re.sub(r'(?<=[:;]) (?=[а-я]\))', r'\n', moduls_doc[key])
    moduls_doc[key] = re.sub(r'(?<=[:;]) (?=[а-яА-Я])(?!\))', r'\n-', moduls_doc[key])
    # Добавление символа новой строки после точки
    moduls_doc[key] = re.sub(r'([а-я]{2,}\)?»?\.)( )(?=[А-Я].*\.)', r'\1\n', moduls_doc[key])

print('-' * 80)
for key in moduls_doc:
    print(f'modul |{key: <20}| from moduls_docs => {moduls_doc[key]}')
    print('-' * 80)

# print(moduls_doc['Функции'])
# print('-' * 80)
# _______________________________________________________________________________________________________________________

# print(moduls_doc['Титульная информация'])
# print('-' * 80)

# Определение ОВУ
ovy_name = re.search(r'(?<=Об\sутверждении\s[Пп]оложения\sо\s)[а-яА-Я -]*(?=\sГенерального\sштаба\sВооруженных\sСил\sРоссийской\sФедерации)|'
                     r'(?<=Об\sутверждении\s[Пп]оложения\sо\s)[а-яА-Я -]*(?=\sВ\sредакции\sприказа)|'
                     r'(?<=ПОЛОЖЕНИЕ\sО\s)[а-яА-Я ]*(?=\sРОССИЙСКОЙ\sФЕДЕРАЦИИ)|'
                     r'(?<=ПОЛОЖЕНИЕ\sо\s)[а-яА-Я- ]*(?=\sГенерального\sштаба\sВооруженных\sСил\sРоссийской\sФедерации)|'
                     r'(?<=[Пп]\s[Оо]\s[Лл]\s[Оо]\s[Жж]\s[Ее]\s[Нн]\s[Ии]\s[Ее]\s[Оо]\s)[а-яА-Я -]*(?=\s(Генерального\sштаба\s)?(Вооруженных\sСил)?\sРоссийской\sФедерации)|'
                     r'(?<=[Пп]\s[Оо]\s[Лл]\s[Оо]\s[Жж]\s[Ее]\s[Нн]\s[Ии]\s[Ее]\s[Оо]\s)[а-яА-Я -]*(?=\s(Министерства\sобороны\sРоссийской\sФедерации))', moduls_doc['Титульная информация'])
# Преобразование слова в родительный падеж
morph = pymorphy2.MorphAnalyzer()
ovy_name = ovy_name[0].split(' ')
# Переменная, в которой сохраняются результаты пробразования
ovy_name_gent = ''
# В цикле проходим по каждому слову
for word in ovy_name:
    # Обработка слова при помощи pymorphy2
    refactor_word = morph.parse(word)[0]
    # Преобразование слова в родительный падеж
    case = refactor_word.inflect({'gent'})
    # Если ovy_name_gent пустая переменная - добавить слово с заглавной буквой,
    # в остальных случаях конкатенация с пробелом
    if ovy_name_gent:
        ovy_name_gent += ' ' + case.word
    else:
        ovy_name_gent += case.word.title()
ovy_name = ovy_name_gent
# print(ovy_name)
# _______________________________________________________________________________________________________________________

# Удаление заглавных строк разделов
for key in moduls_doc:
    if key == 'Полномочия':
        moduls_doc[key] = moduls_doc[key].split('\n')[2:]
        moduls_doc[key] = '\n'.join(moduls_doc[key])
    elif key == 'Задачи':
        moduls_doc[key] = moduls_doc[key].split('\n')[2:]
        moduls_doc[key] = '\n'.join(moduls_doc[key])
    elif key == 'Функции':
        moduls_doc[key] = moduls_doc[key].split('\n')[2:]
        moduls_doc[key] = '\n'.join(moduls_doc[key])
    elif key == 'Общие положения':
        moduls_doc[key] = moduls_doc[key].split('\n')[1:]
        moduls_doc[key] = '\n'.join(moduls_doc[key])
    elif key == 'Руководство':
        moduls_doc[key] = moduls_doc[key].split('\n')[1:]
        moduls_doc[key] = '\n'.join(moduls_doc[key])
    elif key == 'Организация деятельности':
        moduls_doc[key] = moduls_doc[key].split('\n')[1:]
        moduls_doc[key] = '\n'.join(moduls_doc[key])

# Вызов функций, выявляющих отдельные полномочия, задачи, функции и дополняющих их нумерацией
for key in moduls_doc:
    if key == 'Полномочия':
        tasks_Boss = tasks_MO(moduls_doc['Общие положения'])
        for task in tasks_Boss:
            print(f"{task[0]} {task[1]} {task[2]} - {task[3]}")
        print('-' * 80)
        # print(f'modul |{key: <20}| from moduls_docs:')
        moduls_doc['Полномочия'] = authorities(moduls_doc[key])
        for authority in moduls_doc['Полномочия']:
            if len(authority) == 4:
                print(f'{authority[0]} {authority[1]} {authority[2]} - {authority[3]}')
            elif len(authority) == 6:
                print(f'{authority[0]} {authority[1]} абзац {authority[2]} {authority[3]} - <{authority[4]}> {authority[5]}')
        print('Done authorities.')
        # print('-' * 80)
    elif key == 'Задачи':
        print(f'modul |{key: <20}| from moduls_docs:')
        # moduls_doc['Задачи'] = tasks(moduls_doc[key])
        # for task in moduls_doc['Задачи']:
        #     print(f"{task[0]} {task[1]} {task[2]} - {task[3]}")
        # print('-' * 80)
    elif key == 'Функции':
        print(f'modul |{key: <20}| from moduls_docs:')
        # moduls_doc['Функции'] = functions(moduls_doc['Функции'])
        # for function in moduls_doc['Функции']:
        #     if len(function) == 8:
        #         print(f'{function[0]} {function[1]}_{function[2]} {function[3]}| {function[4]} - {function[5]}')
        #     elif len(function) == 10:
        #         print(f'{function[0]} {function[1]}_{function[2]}_{function[3]} {function[4]}| {function[5]} - <{function[6]}> {function[7]}')
        #     elif len(function) == 12:
        #         print(f'{function[0]} {function[1]}_{function[2]}_{function[3]}_{function[4]} {function[5]}| {function[6]} - <{function[7]}> <{function[8]}> {function[9]}')
    #     print('-' * 80)
#     elif key == 'Общие положения':
#         print(f'modul |{key: <20}| from moduls_docs:')
#         print(moduls_doc[key])
#         print('-' * 80)
# _______________________________________________________________________________________________________________________


# _______________________________________________________________________________________________________________________
#
# # Сохранение результатов в формате docx
# document = Document()
#
# # Добавление заголовка. Уровень от 0 до 9 (по умолчанию, 1)
# document.add_heading('Результат работы алгоритма', 0)
# document.add_heading(pdf_file[1], 1)
#
# # Добавление заголовка "Задачи"и соответствующих параграфов
# document.add_heading(f'Задачи {ovy_name}', level=2)
# for paragraph in moduls_doc['Задачи']:
#     document.add_paragraph(f'{paragraph[0]} {paragraph[1]} {paragraph[2]}| {paragraph[3]}')
#
# # Добавление заголовка "Функции" и параграфов с полными текстами функций
# document.add_heading(f'Функции {ovy_name}', level=2)
# for paragraph in moduls_doc['Функции']:
#     if len(paragraph) == 8:
#         document.add_paragraph(f'{paragraph[0]} {paragraph[1]}_{paragraph[2]} {paragraph[3]}| {paragraph[4]} - {paragraph[5]}')
#     elif len(paragraph) == 10:
#         document.add_paragraph(f'{paragraph[0]} {paragraph[1]}_{paragraph[2]}_{paragraph[3]} {paragraph[4]}| {paragraph[5]} - <{paragraph[6]}> {paragraph[7]}')
#     elif len(paragraph) == 11:
#         document.add_paragraph(f'{paragraph[0]} {paragraph[1]}_{paragraph[2]}_{paragraph[3]}_{paragraph[4]} {paragraph[5]}| {paragraph[6]} - <{paragraph[7]}> {paragraph[8]}')
#
# document.save('Форма представления распознанных обязанностей в строку.docx')
# print('Done')
# _______________________________________________________________________________________________________________________
